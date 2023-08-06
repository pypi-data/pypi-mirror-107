from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class President_pr(object):
    def __init__(self):
        self.document = Document()

    def sections(self):
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.5)

    def saver(self, name):
        self.document.save(name)

    def picturer(self, name, width):
        self.document.add_picture(name, width)

    def tabler(self, rows, cols, style):
        self.document.add_table(rows, cols, style)

    def rawer(self, rows, cols, style, spisok):
        table = self.document.add_table(rows, cols, style)
        hdr_cells = table.rows[0].cells
        i = 0
        for f in spisok:
            hdr_cells[i].text = f
            i += 1

    def paragrapher(self, content, font_name='Times New Roman', font_size=12,
                    before_spacing=0, after_spacing=0, line_spacing=1.5,
                    first_line_indent=Inches(0.5), keep_with_next=False):
        paragraph = self.document.add_paragraph(str(content))
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = first_line_indent
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing
        paragraph_format.keep_with_next = keep_with_next

    def run_end(self, content, nur, font_name='Times New Roman', font_size=12, fontrun_bold=True,
                before_spacing=0, after_spacing=0, line_spacing=1.5,
                first_line_indent=Inches(0.5)):
        paragraph = self.document.add_paragraph(str(content))
        run = paragraph.add_run(str(nur))
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        fontrun = run.font
        fontrun.bold = fontrun_bold
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = first_line_indent
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def run_start(self, content, nur, nur_1, font_name='Times New Roman', font_size=12,
                  before_spacing=0, after_spacing=0, line_spacing=1.5,
                  first_line_indent=Inches(0.5)):
        paragraph = self.document.add_paragraph(str(content))
        paragraph.add_run(str(nur)).bold = True
        paragraph.add_run(str(nur_1))
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = first_line_indent
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def paragrapher_otst_up(self, content, font_name='Times New Roman', font_size=12,
                            before_spacing=12, after_spacing=18, line_spacing=1.5,
                            first_line_indent=Inches(0.5)):
        paragraph = self.document.add_paragraph(str(content))
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = first_line_indent
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def paragrapher_prilozhenie(self, content, font_name='Times New Roman', font_size=12,
                                before_spacing=0, after_spacing=0, line_spacing=1.5):
        paragraph = self.document.add_paragraph(str(content))
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def paragrapher_center(self, content, nur, font_name='Times New Roman', font_size=12,
                           before_spacing=24, after_spacing=0, line_spacing=1.5):
        paragraph = self.document.add_paragraph(str(content))
        paragraph.add_run(str(nur)).bold = True
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def paragrapher_center_1(self, content, nur, font_name='Times New Roman', font_size=12,
                             before_spacing=0, after_spacing=36, line_spacing=1.5):
        paragraph = self.document.add_paragraph(str(content))
        paragraph.add_run(str(nur)).bold = True
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def paragrapher_down(self, content, nur, font_name='Times New Roman', font_size=12, fontrun_size=8,
                         before_spacing=306, after_spacing=0, line_spacing=1.5):
        paragraph = self.document.add_paragraph(str(content))
        run = paragraph.add_run(str(nur))
        fontrun = run.font
        fontrun.size = Pt(fontrun_size)
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def paragrapher_down_1(self, content, nur, font_name='Times New Roman', font_size=12, fontrun_size=8,
                           before_spacing=0, after_spacing=0, line_spacing=1.5):
        paragraph = self.document.add_paragraph(str(content))
        run = paragraph.add_run(str(nur))
        fontrun = run.font
        fontrun.size = Pt(fontrun_size)
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing
