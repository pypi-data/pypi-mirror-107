from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def file():
    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.5)

    def paragrapher(content, font_name='Times New Roman', font_size=12,
                    before_spacing=0, after_spacing=0, line_spacing=1.5,
                    first_line_indent=Inches(0.5), keep_with_next=False):
        paragraph = document.add_paragraph(str(content))
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = first_line_indent
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing
        paragraph_format.keep_with_next = keep_with_next

    def run_end(content, nur, font_name='Times New Roman', font_size=12, fontrun_bold=True,
                before_spacing=0, after_spacing=0, line_spacing=1.5,
                first_line_indent=Inches(0.5)):
        paragraph = document.add_paragraph(str(content))
        run = paragraph.add_run(str(nur))
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        fontrun = run.font
        fontrun.bold = fontrun_bold
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = first_line_indent
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def run_start(content, nur, nur_1, font_name='Times New Roman', font_size=12,
                  before_spacing=0, after_spacing=0, line_spacing=1.5,
                  first_line_indent=Inches(0.5)):
        paragraph = document.add_paragraph(str(content))
        paragraph.add_run(str(nur)).bold = True
        paragraph.add_run(str(nur_1))
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = first_line_indent
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def paragrapher_otst_up(content, font_name='Times New Roman', font_size=12,
                            before_spacing=12, after_spacing=18, line_spacing=1.5,
                            first_line_indent=Inches(0.5)):
        paragraph = document.add_paragraph(str(content))
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = first_line_indent
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def paragrapher_prilozhenie(content, font_name='Times New Roman', font_size=12,
                                before_spacing=0, after_spacing=0, line_spacing=1.5):
        paragraph = document.add_paragraph(str(content))
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def paragrapher_center(content, nur, font_name='Times New Roman', font_size=12,
                           before_spacing=24, after_spacing=0, line_spacing=1.5):
        paragraph = document.add_paragraph(str(content))
        paragraph.add_run(str(nur)).bold = True
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def paragrapher_center_1(content, nur, font_name='Times New Roman', font_size=12,
                             before_spacing=0, after_spacing=36, line_spacing=1.5):
        paragraph = document.add_paragraph(str(content))
        paragraph.add_run(str(nur)).bold = True
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def paragrapher_down(content, nur, font_name='Times New Roman', font_size=12, fontrun_size=8,
                         before_spacing=306, after_spacing=0, line_spacing=1.5):
        paragraph = document.add_paragraph(str(content))
        run = paragraph.add_run(str(nur))
        fontrun = run.font
        fontrun.size = Pt(fontrun_size)
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    def paragrapher_down_1(content, nur, font_name='Times New Roman', font_size=12, fontrun_size=8,
                           before_spacing=0, after_spacing=0, line_spacing=1.5):
        paragraph = document.add_paragraph(str(content))
        run = paragraph.add_run(str(nur))
        fontrun = run.font
        fontrun.size = Pt(fontrun_size)
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph_format.line_spacing = line_spacing

    document.add_picture('beta.png', width=Inches(7))
    paragrapher("Аппарат Администрации Смоленской области информирует о начале конкурсного отбора "
                "специалистов для обучения в рамках Государственного плана подготовки "
                "управленческих кадров для организаций народного хозяйства Российской Федерации ("
                "далее - Президентская программа).")
    run_end("Сроки проведения отбора: ", "с 10 марта по 21 мая 2021 года.")
    paragrapher(
        "На обучение приглашаются руководители высшего и среднего звена организаций всех отраслей промышленности, "
        "строительства, транспорта, связи, сферы услуг, торговли, представители малого и среднего бизнеса, "
        "организаций здравоохранения, социальной защиты населения, образования культуры.")
    run_start("", "Требования к конкурсантам: ",
              "возраст до 50 лет, наличие высшего образования, общего стажа работы не "
              "менее "
              "5 лет, опыта работы на управленческих должностях не менее 2 лет, "
              "участи в реализации проекта развития организации.")
    paragrapher("Подготовка специалистов в рамках Президентской программы происходит на базе Смоленского филиала "
                "федерального государственного образовательного бюджетного учреждения высшего образования «Финансовый "
                "университет при Правительстве Российской Федерации». Обучение включает, помимо обучения в российских "
                "образовательных учреждениях, стажировку в ведущих зарубежных организациях на конкурсной основе за "
                "счёт "
                "принимающей стороны.")
    run_start(
        "Заявки на участие в конкурсном отборе по Президенской программе по форме согласно приложению к настоящему "
        "письму предоставляются в Управление государственной гражданской службы и кадровой политики Аппарата "
        "Администрации Смоленской области до ", "11 мая 2021 года ", "по адресу электроной почты: "
                                                                     "gruzdeva_TA@admin-smolensk.ru")
    paragrapher("Подробная информация о Президентской программе размещена в информационно-телекоммуникационной сети "
                "«Интернет» по адресам: ")
    paragrapher("http://apparat.admin-smolensk.ru/prezidentskaya-programma,")
    paragrapher("http://pprog.admin-smolensk.ru.")
    paragrapher("Дополнительную информацию можно получить по телефонам:")
    paragrapher("(4812)38-66-53, 65-98-80, 29-23-47.")
    paragrapher_otst_up("Приложение: на 1 л. в 1 экз.")
    document.add_picture('ateb.png', width=Inches(7))
    paragrapher_down("", "Исп. Груздева Т.А.")
    paragrapher_down_1("", "тел. (4812)29-23-89")
    paragrapher_prilozhenie("Приложение")
    paragrapher_center("", "Заявка на участие в конкурсном отборе")
    paragrapher_center_1("", "по Президентской программе")
    table = document.add_table(rows=2, cols=6, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№ п/п'
    hdr_cells[1].text = 'Ф.И.О.'
    hdr_cells[2].text = 'Место работы'
    hdr_cells[3].text = 'Должность'
    hdr_cells[4].text = 'Контактный телефон'
    hdr_cells[5].text = 'Адрес электронной почты'
    hdr_cells_1 = table.rows[1].cells
    hdr_cells_1[0].text = ' '
    hdr_cells_1[1].text = ' '
    hdr_cells_1[2].text = ' '
    hdr_cells_1[3].text = ' '
    hdr_cells_1[4].text = ' '
    hdr_cells_1[5].text = ' '
    document.save('dem.docx')
