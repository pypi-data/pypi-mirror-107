# Gruppenmanager

import selenium
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.select import Select
from selenium.webdriver.common.keys import Keys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm

import csv, sqlite3
import time, datetime
import random
import re

from iservData import iservObject
from iservToolBox import toolboxObject
from openpyxl import Workbook



class groupsObject:
   
    def __init__(self, iserv):
        """ Legt die beiden Objekte *self.iserv* und *self.tools* an. 
        
            :param Id: (iservObject) Objekt mit dem Rahmendaten für die iServ-Instanz
        """
        self.iserv = iserv
        self.tools = toolboxObject(iserv)



    def create_special_class_groups(self, Id):
        """
        Erstellt Klassengruppen mit bestimmter ID, z.B. Klasse-9a-A
        wird dann als Id = "A" angegeben. Der Klassenlehrer wird als Besitzer
        und Mitglied eingetragen.
        
        :param Id: (str) Bezeichnung, die dem Klassennamen angehängt werden soll.
        """
        self.iserv.login()
        self.iserv.admin_login()
        
        Klassen = self.tools.get_list_of_classes()
        for Klasse in Klassen:
            
            Name = self.tools.get_tutor_name(Klasse)
            Gruppe = "Klasse-" + Klasse + "-" + Id
            self.create_group_with_owner(Gruppe, Name)

        self.iserv.logout()


    def create_special_course_groups(self, Id):
        """
        Erstellt Kursgruppen mit bestimmter ID, z.B. EN92-En-A
        wird dann als Id = "A" angegeben. Der Kurslehrer wird als Besitzer
        und Mitglied eingetragen.
        
        :param Id: (str) Bezeichnung, die dem Klassennamen angehängt werden soll.
        """
        self.iserv.login()
        self.iserv.admin_login()
        
        Klassen = self.tools.get_list_of_courses()
        for Klasse in Klassen:
            print("Kurs: " + Klasse)
            Name = self.tools.get_tutor_name(Klasse)
            Gruppe = self.tools.get_highschool_groupname(Klasse) + "-" + Id
            print("Gruppe:" + Gruppe)
            nr = int(re.findall("\d+", Klasse)[0])
            print(nr)
            
            #self.create_group_with_owner(Gruppe, Name)

        self.iserv.logout()
        
    def create_group_with_owner(self, Gruppenname, Besitzer):
        """Legt eine Gruppe an und legt den Besitzer fest, der dann auch Mitglied der Gruppe ist.

           Erfordert ein vorheriges Login als Admin. Es erfolgt kein Logout.

           :param Gruppenname: (str) Name der Gruppe bei iServ
           :param Gruppenname: (str) Name des Besitzers bei iServ
        """
        
        url = self.iserv.iserv_url + "admin/group/add/"
        self.iserv.browser.get(url)
        time.sleep(8)

        
        input = self.iserv.browser.find_element_by_xpath('//*[@id="group_name"]')
       
        input.send_keys(Gruppenname) 

        field = self.iserv.browser.find_element_by_xpath('//*[@id="group_owner"]')
        select = Select(field)
        try:
            select.select_by_visible_text(Besitzer)
        except:
            self.iserv.message = "Benutzer " + Besitzer + " für Gruppe " + Gruppenname + " nicht gefunden"

        xPath = "/html/body/div[1]/div[2]/div[3]/div[2]/div/div[2]/div/div/div/form/div/div/div/div[3]/table/tbody/tr[2]/td/div/div/input"
        input = self.iserv.browser.find_element_by_xpath(xPath)
        if (input):
            input.send_keys(Besitzer)
        
        xPath='//*[@id="group_actions_submit"]'
        input = self.iserv.browser.find_element_by_xpath(xPath)
        input.click()

    def set_group_owner(self, Gruppenname, Besitzer):
        """Legt den Besitzer einer bestehenden Gruppe fest, der dann auch Mitglied der Gruppe ist.

           Erfordert ein vorheriges Login als Admin. Es erfolgt kein Logout.

           :param Gruppenname: (str) Name der Gruppe bei iServ
           :param Gruppenname: (str) Name des Besitzers bei iServ
        """
        
        url = self.iserv.iserv_url + "admin/group/edit/" + Gruppenname.lower()
        self.iserv.browser.get(url)
        time.sleep(4)

        field = self.iserv.browser.find_element_by_xpath('//*[@id="group_owner"]')
        select = Select(field)
        try:
            select.select_by_visible_text(Besitzer)
        except:
            self.iserv.message = "Benutzer " + Besitzer + " für Gruppe " + Gruppenname + " nicht gefunden"
            
        xpath = '//*[@id="group_actions_submit"]'
        
        input = self.iserv.browser.find_element_by_xpath(xpath)
        input.click()
      
    def set_tutors_as_group_owners(self):
        """Setze die Klassenlehrer als Besitzer ihrer Klassengruppen ein."""
        Klassen = self.tools.get_list_of_classes()
        self.iserv.login()
        self.iserv.admin_login()
        for Klasse in Klassen:
            Gruppe = self.tools.get_class_groupname(Klasse)
            Lehrer = self.tools.get_tutor_name(Klasse)
            self.set_group_owner(Gruppe, Lehrer)
        self.iserv.logout()
        self.iserv.message ="Klassenlehrer als Besitzer der Klassengruppen eingetragen"
      
    def update_teachers_in_class_groups(self):
        """Aktualisiert die Mitgliedschaften von Lehrern in den Klassengruppen.
        
           - Liest die Lehrer aus den iServ Klassengruppen aus der Datenbank aus.
           - Liest die Lehrer der Klassen aus dem Stundenplan.
           - Übersetzt die Kürzel aus dem Stundenplan in ganze Namen.
           - Vergleicht die Namenslisten aus iServ und dem Stundenplan
           - Korrigiert die Angaben in den iServ-Listen online
           
        """
        conn = sqlite3.connect(self.iserv.database)
        c = conn.cursor()
    
        Klassen = self.tools.get_list_of_classes()
        
        self.iserv.login()
        self.iserv.admin_login()
        
        for Klasse in Klassen:
            Gruppe = "klasse-" + Klasse.lower()
            sql = "SELECT Vorname, Nachname, Rolle FROM group_members_iserv WHERE Gruppe = ? AND Rolle = ?"
            c.execute(sql, (Gruppe, "Lehrer"))
            
            Lehrer = c.fetchall()
            Lehrer_iserv = []
            for L in Lehrer:
                Lehrer_iserv.append(L[0] + " " + L[1])
            
            sql = "SELECT DISTINCT Lehrer FROM Stundenplan WHERE Klasse = ?"
            c.execute(sql, (Klasse,))
            
            Lehrer = c.fetchall()
            Lehrer_timetable = []
            
            for L in Lehrer:
                sql = "SELECT Vorname, Name FROM Lehrer WHERE Kurz = ?"
                c.execute(sql, (L[0],))
                Name = c.fetchall()
                if len(Name) > 0:
                    Lehrername = Name[0][0] + " " + Name[0][1]
                else:
                    Lehrername = "N.N."
                Lehrer_timetable.append(Lehrername)
                
                if (Lehrername not in Lehrer_iserv and Lehrername != "N.N."):
                    pass
                    #print("nicht in der Liste der Klasse " + Klasse + " : "+ Lehrername)
            
            
            for L in Lehrer_iserv:
                if L not in Lehrer_timetable:
                    self.delete_member_from_class(L, Klasse)
                    
        self.iserv.logout()
        
        c.close
        conn.close
        
    def delete_member_from_class(self, Name, Klasse):
        """Löscht ein Mitglied in einer Gruppe"""
        print("delete " + Name + " from " + Klasse)
        
    
    def read_users_grade(self, Jahr):
        """List eine Liste aller aktiven User in die Datenbank ein, um den Namen Benutzernamen zuordnen zu können.
        
           Dieser Vorgang dauert einige Minuten.
    
           table: active_users
           - Vorname
           - Nachname
           - Account
        """
        conn = sqlite3.connect(self.iserv.database)
        c = conn.cursor()
        c.execute('DROP TABLE IF EXISTS active_users')
        c.execute('CREATE TABLE active_users (Account TEXT, Vorname TEXT, Nachname TEXT)')
        
        url = self.iserv.iserv_url + "admin/user?filter%5Bstatus%5D=active&filter%5Bgroups%5D%5B%5D=jahrgang-"+ str(Jahr) + "&filter%5Bsearch%5D="
        self.iserv.browser.get(url)
        time.sleep(8)
        
        xpath_table = '/html/body/div[1]/div[2]/div[3]/div[2]/div/div[2]/div/div/form/div[1]/table'
        
        Zeilen = self.iserv.browser.find_elements_by_xpath(xpath_table + "/tbody[1]/tr")
        i = 0
        for row in Zeilen:
            Zeile = row.find_elements_by_tag_name("td")
       
            if (len(Zeile) > 1):
                i = i + 1
                Account = Zeile[1].text
                Vorname = Zeile[2].text
                Nachname = Zeile[3].text
                
                c.execute("INSERT INTO active_users (Account, Vorname, Nachname) VALUES (?,?,?)", (Account, Vorname, Nachname))
                conn.commit()   
                if (i % 100 == 0):
                    print(i)

        c.close
        conn.close
        self.iserv.write_table_timestamp("active_users")
        
    def get_account_by_name(self, Vorname, Nachname):
        """Ermittelt den Account eines Benutzers anhand des Vor- und Nachnamens
        
           :param Vorname: (str) Vorname
           :param Nachname: (str) Nachname
        """
        conn = sqlite3.connect(self.iserv.database)
        c = conn.cursor()
        
        sql = "SELECT Account FROM active_users WHERE Vorname = ? AND Nachname = ?"
        c.execute(sql, (Vorname, Nachname))
            
        result = c.fetchall()
        
        if (len(result) > 0):
            account = result[0][0]
        else:
            account =""
        
        c.close
        conn.close
        
        return account
        
    def read_members_of_class_groups(self):
        """Liest die Mitglieder aller Klassengruppen in die Datenbank ein.
        
        """
        Liste = self.tools.get_list_of_classes()
        print("Lese Mitglieder der Gruppe der Klassen")
        
        self.iserv.login()
        self.iserv.admin_login()
        
        for Klasse in Liste:
            group_name = "klasse-" + Klasse.lower()
            self.read_class_group_members(group_name)
            print("... " + group_name)
            
        self.iserv.logout()
    
    def read_class_group_members(self, Gruppenname):
        """Liest die Mitglieder einer Klassengruppe aus iServ aus und trägt sie in die Datenbank ein. 
        
           table: group_members_iserv
           
           - Gruppe 
           - Vorname
           - Nachname
           - Rolle
           - Pos
        
           :param Gruppenname: (str) Name der Gruppe bei iServ, z.B. "klasse-9fa"
           
        """
        
        time.sleep(2)
        url = self.iserv.iserv_url + "admin/group/edit/" + Gruppenname 
    
        self.iserv.browser.get(url)
        time.sleep(8)
        link = self.iserv.browser.find_element_by_partial_link_text("Mitglieder").click()
        time.sleep(8)
        xpath_table = '/html/body/div/div[2]/div[3]/div[2]/div/div/div/form/div[2]/div/div/div[2]/div[1]/table'
        
        css_table =".table"
        
        Zeilen = self.iserv.browser.find_elements_by_xpath(xpath_table + "/tbody[1]/tr")
        rows_max = len(Zeilen)

        i = 0      
        
        conn = sqlite3.connect(self.iserv.database)
        c = conn.cursor()
        
        c.execute('CREATE TABLE IF NOT EXISTS group_members_iserv (Gruppe TEXT, Vorname TEXT, Nachname TEXT, Rolle TEXT, Pos INTEGER)')
 
        c.execute ("DELETE FROM group_members_iserv WHERE Gruppe = ?", (Gruppenname,))
        
        for row in Zeilen:
            Zeile = row.find_elements_by_tag_name("td")
            i += 1
            if (len(Zeile) > 1):
                
                Vorname = Zeile[0].text
                Nachname = Zeile[1].text
                Klasse = Zeile[2].text
                Rolle = Zeile[3].text
                Nummer = i             
           
                if ("Lehrer" in Rolle):
                    Rolle = "Lehrer"
                else:
                    Rolle = "Schüler"
        
                c.execute("INSERT INTO group_members_iserv (Gruppe, Vorname, Nachname, Rolle, Pos) VALUES (?,?,?,?,?)", (Gruppenname, Vorname, Nachname, Rolle, Nummer))
                conn.commit()   
                
        c.close
        conn.close
    
        
    def delete_group_members(self, Gruppenname, Liste):
        """Löscht eine Liste von Mitlgiedern aus einer Gruppe
        
        :param Gruppenname: (str) Name der Gruppe bei iServ, z.B. "klasse-9fa"
        :param Liste: (list) Liste der Namen der zu löschenden Mitlglieder
        
        """
        self.login()
        self.admin_login()
        
        time.sleep(2)
        url = self.iserv_url + "admin/group/edit/" + Gruppenname 
    
        self.browser.get(url)
        time.sleep(8)
        link = self.browser.find_element_by_partial_link_text("Mitglieder").click()
        time.sleep(8)
        xpath_table = '/html/body/div/div[2]/div[3]/div[2]/div/div/div/form/div[2]/div/div/div[2]/div[1]/table'
        
        css_table =".table"
        
        Zeilen = self.iserv.browser.find_elements_by_xpath(xpath_table + "/tbody[1]/tr")
        rows_max = len(Zeilen)

        i = 0      
        
        for row in Zeilen:
            Zeile = row.find_elements_by_tag_name("td")
            i += 1
            if (len(Zeile) > 1):
                
                Vorname = Zeile[0].text
                Nachname = Zeile[1].text
                Klasse = Zeile[2].text
                Rolle = Zeile[3].text      
                Nummer = i             
                print(str(i) + "\t" + Nachname + ", " + Vorname ) 
                
        xpath = '//*[@id="group_actions_submit"]'
        Eintrag = self.browser.find_element_by_xpath(xpath)
        Eintrag.click()
        
        self.logout()
        
    def set_user_password(self, user_name, password):
        """Setzt ein Passwort für einen iServ-User
        
           Erfordert ein vorheriges Login als Admin. Es erfolgt kein Logout.

           :param user_name: (str) Login-Name bei iServ
           :param password: (str) zu setzendes Passwort
        """
        url = self.iserv.iserv_url + "admin/user/password/" + user_name
        self.iserv.browser.get(url)
        time.sleep(2)
        input = self.iserv.browser.find_element_by_name("passwordForm[password]")
        input.send_keys(password)
        
        xpath = '//*[@id="passwordForm_actions_approve"]'
        setzen = self.iserv.browser.find_element_by_xpath(xpath)
        setzen.click()
        
    def generate_user_password(self):
        """Generiert ein merkbares Passwort und gibt es zurück.
          
           Das Passwort wird aus einer Farbe, einem Nomen, einem Sonderzeichen und einer zweistelligen Zahl zusammengesetzt, z.B. *gelbWolke%82*
        """
        nouns = ["Wolf", "Tiger", "Haus", "Buch", "Auto", "Maus", "Blume", "Hand","Baum", "Gras", "Wolke", "Roller", "Rose", "Schaf", "Salz", "Garten", "Topf", "Kiste", "Fisch", "Vogel", "Katze", "Busch","Wald","Haus","Fisch","Tasche","Zahn"]
        colors = ["gelb", "rot", "blau", "braun", "orange", "lila", "pink", "schwarz","klein","rund", "eckig","leicht","schwer","kariert","alt","neu"]
        special = ["?", "%", "!", "#", "&", "*", "/"]
        noun_id = random.randrange(len(nouns))
        color_id = random.randrange(len(colors))
        special_id = random.randrange(len(special))
        number = str(random.randrange(11, 99))
        password = colors[color_id] + nouns[noun_id] + special[special_id] + number
        return password
    
    def generate_passwords_for_class(self, Klasse):
        """Generiert neue Passwörter für eine Klasse, setzt sie in iServ und generiert druckbare Hinweise für Schüler als    Word-Dokument.
        
           :param Klasse: (str) Name der Klasse
        """
        
        Klassengruppe = "klasse-" + Klasse.lower()
        #Gruppen.read_class_group_members(Klassengruppe)
        Klassenlehrer = self.tools.get_tutor_name(Klasse)
        
        conn = sqlite3.connect(self.iserv.database)
        c = conn.cursor()
        
        c.execute("SELECT Vorname, Name FROM Schueler WHERE Klasse = ? COLLATE NOCASE", (Klasse,))
        Data = c.fetchall()
        
        doc = Document()
        
        section = doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        
        doc.add_heading('Passwörter für Klasse ' + Klasse, 0)
        doc.add_paragraph("")
        doc.add_heading("Klassenlehrer: " + Klassenlehrer, 1)
        Klassenlehrer_Mail = self.tools.get_account_by_name(Klassenlehrer) + self.iserv.mail_domain
        
        doc.add_paragraph("")
        doc.add_heading(Klassenlehrer_Mail, 2)
        
        for row in Data:
            Vorname = row[0]
            Nachname = row[1]
        
            password = Gruppen.generate_user_password()
            account = Gruppen.get_account_by_name(Vorname, Nachname)
            Gruppen.set_user_password(account, password)
        
            self.print_info_page(Vorname, Nachname, account, password, Klasse, Klassenlehrer, Klassenlehrer_Mail, doc)
        
        
        doc.save(self.iserv.data_dir + "/iserv-Logins-" + Klasse + ".docx")
        
        c.close
        conn.close
        
    def generate_passwords_for_year(self, Jahrgang):
        """Generiert Passwörter für einen Jahrgang
        
           :param Jahrgang: (int) Jahrgang, für den die Passwörter generiert und gesetzt werden sollen.
        """
        Klassen = self.tools.get_classes_in_year(Jahrgang)
        for Klasse in Klassen:
            self.generate_passwords_for_class(Klasse)
            
            
    def generate_passwords_for_guests(self):
        """Generiert Passwörter für Gastzugänge"""
        
        
        if (2 > 1):
            Kurz = "Ed"
            Name = self.tools.get_teacher_name_by_short(Kurz)
            Account = self.tools.get_account_by_name(Name)
            Passwort = self.generate_user_password()
            #self.groups.set_user_password(Account, Passwort)
            print()
            print(Name)
            print(Account)
            print(Passwort)
        
        
    def print_info_page(self, Vorname, Nachname, account, password, Klasse, Klassenlehrer, Lehrer_Mail, doc):
        """Schreibt die Informationen für einen Schüler auf eine Word-Seite
        
           :param Vorname: (str) 
           :param Nachname: (str)
           :param account: (str) iServ-Account, z.B. fritz.fischer
           :param password: (str)
           :param Klasse: (str)
           :param Klassenlehrer: (str) Name des Klassenlehrers
           :param doc: (docx) Bezug zum Word-Dokument
        
        """
        doc.add_page_break()
        if (Klasse != ""):
            p = doc.add_paragraph("Klasse " + Klasse)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        doc.add_heading('iServ-Zugang für  ' + Vorname + " " + Nachname, 0)
        
        info = "Du findest unseren Schulserver iServ unter dieser Adresse:"
        doc.add_paragraph(info)
        doc.add_heading(self.iserv.iserv_url, level=3)
        #doc.add_paragraph("")
        info = "Du kannst dich mit dem folgenden Benutzernamen anmelden:"
        doc.add_paragraph(info)
      
        doc.add_heading(account, level=3)
        
        info = "\nBeachte beim Schreiben deines Namens bitte Folgendes:"
        doc.add_paragraph(info)
        doc.add_paragraph('Dein Benutzername besteht aus deinem Vor- und Nachnamen.', style='List Bullet')
        doc.add_paragraph('Zwischen Vor- und Nachnamen steht ein Punkt - KEIN Leerzeichen.', style='List Bullet')
        doc.add_paragraph('Schreibe den Namen bitte vollständig klein.', style='List Bullet')
        
        doc.add_paragraph('Beachte, dass es beim Benutzernamen keine Umlaute (Ä, Ö, Ü) und kein ß gibt. Aus einem "ä" wird "ae", aus einem "ß" wird "ss".', style='List Bullet')
        doc.add_paragraph('Akzente entfallen. So wird aus "André" ein "andre".', style='List Bullet')
        #doc.add_paragraph("")
        doc.add_paragraph("Dein Passwort lautet:")
        doc.add_heading(password, level=3)
        
        info = "\nBeachte bei deinem Passwort bitte Folgendes:"
        doc.add_paragraph(info)
        doc.add_paragraph('Achte genau auf die Groß-/Kleinschreibung.', style='List Bullet')
        doc.add_paragraph('Du findest die Sonderzeichen bei den Zahlen.', style='List Bullet')
        doc.add_paragraph('Du kannst dein Passwort später ändern.', style='List Bullet')
        
        #doc.add_paragraph("")
        doc.add_heading('Bewahre dein Passwort so auf, dass nur du selbst darauf Zugriff hast. - Gib es nie weiter, auch nicht an beste Freunde!', level=3)
        
        #doc.add_paragraph("")
        doc.add_paragraph('Falls eine andere Person mit deinem Login etwas schreibt oder tut, geschieht das in DEINEM Namen. Du bist verantwortlich für alles, was in deinem Namen geschieht', style='List Bullet')
        doc.add_paragraph('Wenn du Probleme mit deinem Passwort hast, wende dich bitte an Frau Rippe im Schülerbüro.', style='List Bullet')
        doc.add_paragraph('Solltest du andere Probleme mit iServ haben, wende dich an deinen Klassenlehrer bzw. deine Klassenlehrerin.', style='List Bullet')
        doc.add_paragraph('Auf der Homepage https://www.c-a-g.de findest du unter "ISERV/iServ allgemein" Informationen zu unserem iServ', style='List Bullet')
        if (Lehrer_Mail != ""):
            doc.add_heading("Sobald du das geschafft hast, schreibe eine Mail an " + Lehrer_Mail + " ", level=3)
           
    def create_highschool_courses(self):
        """Legt die Oberstufenkurse bei iServ an und legt die Kursleiter als Besitzer fest."""
        
        conn = sqlite3.connect(self.iserv.database)
        c = conn.cursor()
        
        c.execute("SELECT Gruppe, Kursleitung FROM Kurse")
        Gruppen = c.fetchall()
        
        self.iserv.login()
        self.iserv.admin_login()
        if (Gruppen):
            for row in Gruppen:
                Kurs = row[0]
                Kursleitung = row[1]
                self.create_group_with_owner(Kurs, Kursleitung)
                
        self.iserv.logout()
        
        c.close
        conn.close
        
   ### 
    def import_groups_from_list(self):
        """Legt neue Gruppen nach einer Liste an.
           
           Dazu muss sich im Datenverzeichnis eine Datei **Gruppen-neu.csv** mit dem Schema *[Lehrerkürzel] ; [Name der Gruppe]* befinden. 
        
        """
        with open(self.iserv.data_dir + "/Gruppen-neu.csv", 'rt', encoding="utf-8") as csvfile:
            Gruppen = csv.reader(csvfile, delimiter=",")
            
            self.iserv.login()
            self.iserv.admin_login()
            
            for row in Gruppen:
                Name = row[6].strip()
                Name = Name.replace(";","")
                Name = Name.replace(",","")
                Name = Name.replace("/","-")
                
                Gruppe = Name + "-" + str(self.iserv.end_of_year)
                Lehrer = self.tools.get_teacher_name_by_short(row[5].strip())
                self.create_group_with_owner(Gruppe, Lehrer)
                
            self.iserv.logout()
                
        self.iserv.message = "Gruppen angelegt."
            
    def set_group_as_mailing_list(self, Gruppe):
        """Setze Gruppe als Mailing-Liste, d.h. entferne den Dateibereich, das Forum und den Kalender
        
            :param Gruppe: (str) Name der zu bearbeitenden Gruppe
        """
        self.iserv.login()
        self.iserv.admin_login()
        url = self.iserv.iserv_url + "admin/group/edit/" + Gruppe.lower()
        self.iserv.browser.get(url)
        time.sleep(4)
        xpath = "/html/body/div[1]/div[2]/div[3]/div[2]/div/div[2]/div/div/div/form/div[2]/div/div/div[1]/div[5]/div/span/span[1]/span/ul/li[1]/span[2]"
        xpath="/html/body/span/span"
        link = self.iserv.browser.find_element_by_xpath(xpath)
        link.click()
        
        xpath = '//*[@id="select2-group_flags-result-zrxk-has_home"]'
        link = self.iserv.browser.find_element_by_xpath(xpath)
        link.click()
            ## 
            
    def get_teachers_in_subject_of_year(self, Fach, Jahr):
        """Ermittelt alle Lehrer, die in einem Jahr in einem Fach unterrichten."""
        self.iserv.message = "Ermittle alle Lehrer im Fach " + Fach + " im Jahrgang " + str(Jahr)
        conn = sqlite3.connect(self.iserv.database)
        c = conn.cursor()
        
        c.execute("SELECT Lehrer_kurz, Klasse FROM Lehrer_in_Klasse WHERE Fach = ? and Jahrgang = ?", (Fach, Jahr))
        data = c.fetchall()
        
        c.close
        conn.close
        
        return data
    
    
    def get_list_of_teachers_in_subject(self, Fach):
        """Erstelle eine Liste der Lehrer in einem Fach"""
        doc = Document()
        Jahre = [5,6,7,8,9,10,11]
        for Jahr in Jahre:
            data = self.get_teachers_in_subject_of_year(Fach, Jahr)
            if (data):
                doc.add_paragraph(Fach + " Jahrgang " + str(Jahr))
                for row in data:
                    Lehrer = row[0]
                    Klasse = row[1]
                    Name = self.tools.get_teacher_name_by_short(Lehrer)
                    doc.add_paragraph(Klasse + "    \t" + Lehrer + "\t" + Name , style='List Bullet')
                    
        doc.save(self.iserv.data_dir + "/Lehrer-im-Fach-" + Fach + ".docx")
       
    def export_list_of_teachers_in_all_subjects(self):
        """Ermittelt die Lehrer in einem Fach jahrgangsweise"""
        Liste = self.tools.get_list_of_all_subjects()
        for Fach in Liste:
            self.get_list_of_teachers_in_subject(Fach)
            
        
        
    def export_list_of_students_in_class(self, Klasse):
        """Exportiert eine Schülerliste der Klasse als Word-Dokument"""
        conn = sqlite3.connect(self.iserv.database)
        c = conn.cursor()
        
        c.execute("SELECT Name, Vorname, Id FROM Schueler WHERE Klasse = ? ORDER BY Name",(Klasse,))
        data = c.fetchall()
        
        if (data):
            doc = Document()
            wb = Workbook()
            ws = wb.active
            section = doc.sections[0]
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            
            doc.add_heading("Klasse " + Klasse, level=0)
            datei = open(self.iserv.data_dir + "/export/Klassen/Klasse-" + Klasse + ".csv",'w')
            datei.write("Name, Vorname\n")
            i = 1
            ws.cell(column=2, row=1, value="Name")
            ws.cell(column=1, row=1, value="Vorname")
            #ws.cell(column=3, row=1, value="Name")
            #ws.cell(column=4, row=1, value="Vorname")
            #ws.cell(column=5, row=1, value="Username")
            
            for row in data:
                Name = row[0]
                Vorname = row[1]
                Id = row[2]
                datei.write(Name + "," + Vorname + "\n")
                doc.add_paragraph(Name + ", " + Vorname , style='List Bullet')
                i += 1
                """
                anonName = str(Id)
                anonName = Vorname[0] + anonName[9:16]
                anonVorname = Vorname[0:2] + anonName[-2] + anonName[-1]
                """
                ws.cell(column=2, row=i, value=Name)
                ws.cell(column=1, row=i, value=Vorname)
                #ws.cell(column=3, row=i, value=anonName)
                #ws.cell(column=4, row=i, value=anonVorname)
                #ws.cell(column=5, row=i, value=anonName)
    
            doc.save(self.iserv.data_dir + "/export/Klassen/Klasse-" + Klasse + ".docx")
            datei.close()
            wb.save(self.iserv.data_dir + "/export/Klassen/Klasse-" + Klasse + ".xlsx")
        c.close
        conn.close
        
            
    
        
    def export_list_of_teachers_in_class(self, Klasse):
        """Exportiert eine Liste der Lehrer in einer Klasse als Word-Dokument"""
        conn = sqlite3.connect(self.iserv.database)
        c = conn.cursor()
        
        c.execute("SELECT Lehrer_kurz, Fach FROM Lehrer_in_Klasse WHERE Klasse = ? COLLATE NOCASE ORDER BY Fach",(Klasse,))
        data = c.fetchall()
        if (data):
            doc = Document()
            section = doc.sections[0]
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            
            doc.add_heading("Klasse " + Klasse, level=0)
            Klassenlehrer = self.tools.get_tutor_name(Klasse)
            doc.add_heading("Klassenlehrer: " + Klassenlehrer + "\n\n", level=2)
            
            for row in data:
                Kurz = row[0]
                Fach = row[1]
                Name = self.tools.get_teacher_name_by_short(Kurz)
                doc.add_paragraph(Fach + "     \t" + Kurz + "    \t" + Name, style='List Bullet')
    
            doc.save(self.iserv.data_dir + "/Lehrer-in-Klasse-" + Klasse + ".docx")
        c.close
        conn.close
        
        
    def export_list_of_teachers_in_all_classes(self):  
        """Exportiert die Listen der Fachlehrer in allen Klassen."""
        Klassen = self.tools.get_list_of_classes()
        self.iserv.message =" Exportiere Listen der Fachlehrer in den Klassen"
        for Klasse in Klassen:
            self.export_list_of_teachers_in_class(Klasse)
        
        
    def export_list_of_students_in_all_classes(self):
        """Exportiert die Schülerlisten aller Klassen."""
        Klassen = self.tools.get_list_of_classes()
        for Klasse in Klassen:
            self.export_list_of_students_in_class(Klasse)
        
            
    def set_group_as_email_list(self, Gruppe):
        """Setze die Gruppe als eMail Liste, d.h. lösche den Dateibereich, das Forum und den Kalender einer Gruppe. Dies trägt zur Übersichtlichkeit bei."""
        self.iserv.message = ""
        url = self.iserv.iserv_url + "admin/group/edit/" + Gruppe.lower()
        self.iserv.browser.get(url)
        time.sleep(6)
        
     
        field = self.iserv.browser.find_element_by_name("group[flags][]")
        select = Select(field)

        try:
            select.deselect_by_visible_text("Gruppe hat Dateibereich")
        except:
            self.iserv.message = "Fehler"

        try:
            select.deselect_by_visible_text("Gruppe hat Forum")
        except:
            self.iserv.message = "Fehler"
 
        try:
            select.deselect_by_visible_text("Gruppe hat Kalender")
        except:
            self.iserv.message = "Fehler"
            
        xpath = '//*[@id="group_actions_submit"]'
        link = self.iserv.browser.find_element_by_xpath(xpath)
        link.click()
            
    def export_iserv_list_parent_reps(self):
        """Exportiert eine Liste mit Accounts für Elternvertreter für den Import in iServ
        """
        Liste = self.tools.get_list_of_classes()
        print("Lese Mitglieder der Gruppe der Klassen")
        Infotext  = "\n\n\n\niServ-Login  für Elternvertreter/innen\n"
        Infotext += self.iserv.iserv_url + "\n\n"
        Infotext += "Benutzername / Passwort\n"
        
        Reps = open(self.iserv.data_dir + "/iserv-parent-reps.csv", "w")
        
        for Klasse in Liste:
            KL = open(self.iserv.data_dir + "/Elternvertreter-" + Klasse + ".txt", "w")
            Nachname = "Klasse-" + Klasse + "-" + str(self.iserv.end_of_year)
            Jahrgang = self.tools.get_year_by_class(Klasse)
            Gruppen  = "Klasse-" + Klasse + "-" + str(self.iserv.end_of_year) + "-EV,"
            Gruppen += "Elternvertreter-"  + str(self.iserv.end_of_year) 
            Gruppen += ",Elternvertreter-"  + str(Jahrgang) + "er-" + str(self.iserv.end_of_year)
            for i in ["3","2","1"]:
                Passwort = self.generate_user_password()
                Vorname = "Eltern-" + i
                KL.write(Infotext)
                KL.write(Vorname.lower() + "." + Nachname.lower() + "\n" + Passwort + "\n\n\n\n\n\n\n\n")
                if (i == "1"):
                    Gruppen += ",EV-Vorsitzende-" + str(self.iserv.end_of_year)
                Reps.write('\"' + Vorname + '\",\"' + Nachname + '\",\"' + Passwort + '\",\"' + Gruppen + '\"\n')
            KL.close()
        Reps.close()
            
            
            

###########################




if __name__ == '__main__':
    print()
    print()

    print("GRUPPEN")
    print()

    iServ = iservObject()
    Gruppen = groupsObject(iServ)
    
    #Gruppen.generate_passwords_for_guests()
    Gruppen.export_iserv_list_parent_reps()
    """
    
    Gruppen.iserv.login()
    Gruppen.iserv.admin_login()
    Gruppen.create_special_course_groups("A")
    for Klasse in Klassen:
        Gruppe = "Klasse-" + Klasse.lower() + "-2021-"
        
        Gruppen.set_group_as_email_list(Gruppe + "l")
        Gruppen.set_group_as_email_list(Gruppe + "s")
        

    
    print(Gruppen.iserv.message)
    """